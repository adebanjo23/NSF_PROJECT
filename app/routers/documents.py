from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import List
import uuid
from datetime import datetime

from app.database import get_db
from app.models import User, Document
from app.auth import get_current_user, require_admin
from app.s3_service import s3_service
from app.graphrag_service import graphrag_service

router = APIRouter()


class DocumentResponse(BaseModel):
    id: int
    filename: str
    processed: bool
    uploaded_at: datetime
    file_size: int


class UploadResponse(BaseModel):
    message: str
    document_id: int


@router.post("/upload", response_model=UploadResponse)
async def upload_document(
        file: UploadFile = File(...),
        current_user: User = Depends(get_current_user),
        db: Session = Depends(get_db)
):
    if current_user.role not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")

    if not file.filename.lower().endswith(('.pdf', '.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

    file_key = f"documents/{uuid.uuid4()}_{file.filename}"

    if not s3_service.upload_file(file.file, file_key):
        raise HTTPException(status_code=500, detail="Failed to upload file")

    document = Document(
        filename=file.filename,
        s3_key=file_key,
        uploaded_by=current_user.id,
        file_size=file.size if hasattr(file, 'size') else 0,
        content_type=file.content_type
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    return UploadResponse(
        message="File uploaded successfully",
        document_id=document.id
    )


@router.get("/", response_model=List[DocumentResponse])
def list_documents(
        current_user: User = Depends(get_current_user),
        db: Session = Depends(get_db)
):
    documents = db.query(Document).order_by(Document.uploaded_at.desc()).all()
    return documents


@router.post("/process/{document_id}")
def process_document(
        document_id: int,
        admin_user: User = Depends(require_admin),
        db: Session = Depends(get_db)
):
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    if document.processed:
        raise HTTPException(status_code=400, detail="Document already processed")

    file_content = s3_service.download_file(document.s3_key)
    if not file_content:
        raise HTTPException(status_code=500, detail="Failed to download file from S3")

    try:
        if document.filename.lower().endswith('.pdf'):
            text_content = extract_pdf_text(file_content)
        elif document.filename.lower().endswith(('.docx', '.doc')):
            text_content = extract_docx_text(file_content)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        document_content = f"""
DOCUMENT METADATA:
- Filename: {document.filename}
- Uploaded: {document.uploaded_at}
- File Size: {document.file_size} bytes

DOCUMENT CONTENT:
{text_content}
"""

        graphrag_service.add_document(document_content)

        document.processed = True
        db.commit()

        return {"message": "Document processed successfully"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")


@router.delete("/{document_id}")
def delete_document(
        document_id: int,
        admin_user: User = Depends(require_admin),
        db: Session = Depends(get_db)
):
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    s3_service.delete_file(document.s3_key)
    db.delete(document)
    db.commit()

    return {"message": "Document deleted successfully"}


def extract_pdf_text(content: bytes) -> str:
    import PyPDF2
    import io

    pdf_file = io.BytesIO(content)
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"

    return text.strip()


def extract_docx_text(content: bytes) -> str:
    from docx import Document
    import io

    docx_file = io.BytesIO(content)
    doc = Document(docx_file)

    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    return "\n".join(text_parts)